"""
Build the Word workbook for the Nomade Vans case — aligned to Nomade v3.0
(4 formal questions, strategic 4-scenario framework, competitor attributes,
daily-only competitor data disclaimer). The interactive simulator covers the
experimentation; this workbook is the paper reference.
"""
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SIMULATOR_URL = "https://cmoreno34.github.io/nomade-pricing-simulator/"

ORANGE = RGBColor(0xF9, 0x73, 0x16)
BLUE   = RGBColor(0x25, 0x63, 0xEB)
GREY   = RGBColor(0x64, 0x74, 0x8B)
BLACK  = RGBColor(0x0F, 0x17, 0x2A)
AMBER  = RGBColor(0x92, 0x40, 0x0E)


def set_cell_bg(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_fill)
    tc_pr.append(shd)


def h(doc, text, level=1, color=BLACK):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.color.rgb = color
    return p


def para(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)
    return p


def note_box(doc, text, fill='FEF3C7', title='⚠ Data note', title_color=AMBER):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(title + '  ')
    r.bold = True; r.font.color.rgb = title_color
    p.add_run(text)
    doc.add_paragraph()


def answer_box(doc, lines=10):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, 'F8FAFC')
    p = cell.paragraphs[0]
    p.add_run('\n' * (lines - 1))
    doc.add_paragraph()


def make_header_row(tbl, headers):
    for i, txt in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = txt
        for r in cell.paragraphs[0].runs:
            r.bold = True


def build():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ---------- Cover ----------
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('INTERACTIVE PRICING STRATEGY — CASE STUDY')
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = GREY

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('Nomade Vans — Pricing under Competition and Psychological Factors')
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = BLUE

    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('Simulator-driven edition · aligned to Nomade v3.0 — student workbook')
    r.italic = True; r.font.size = Pt(12); r.font.color.rgb = GREY

    doc.add_paragraph()
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('© César Moreno Pascual PhD — based on Nomade v2.0 / v3.0\n').italic = True
    meta.add_run('Interactive simulator: ').italic = True
    link = meta.add_run(SIMULATOR_URL); link.font.color.rgb = BLUE; link.underline = True

    doc.add_page_break()

    # ---------- 1. The story ----------
    h(doc, '1. The case — Nomade Vans')
    para(doc,
         "Nomade Vans is a Spanish company focused on the design and "
         "\"camperization\" of sustainable vans. Customers configure their "
         "vehicle through the company website and the firm commits to a "
         "one-month delivery. Unlike most players in the category, Nomade "
         "uses a unique design aligned with sustainability values and the "
         "natural environment.")
    para(doc,
         "After two seasons designing and camperizing vans for direct sale, "
         "the founders are now considering a new line of business: "
         "renting camperized vans. To launch it, the company has defined "
         "two service tiers, three rental periods and a survey-based "
         "willingness-to-pay analysis.")

    para(doc, 'Two service tiers:', bold=True)
    bullet(doc, 'Standard Camper Rental — base conversion with proven kitchen + bed package.')
    bullet(doc, 'Premium Camper Rental — hybrid model, extra bed, indoor shower, air conditioning, solar power, extra kitchenware.')

    para(doc, 'Three rental periods:', bold=True)
    bullet(doc, 'Day — single-day rental.')
    bullet(doc, 'Weekend — 2 nights.')
    bullet(doc, 'Week — 6 nights.')
    para(doc,
         'The combination of tiers × periods defines a 2 × 3 price structure.',
         italic=True, color=GREY)

    para(doc,
         "To calibrate willingness-to-pay, Nomade surveyed 95 people, asking "
         "about demographics, travel habits and the maximum price they would "
         "be willing to pay in each of the six rental scenarios. The marketing "
         "department also identified three competitors in the relevant market — "
         "Further VAN, People Camper and Ocean Vans — currently publishing "
         "daily tariffs only.")

    note_box(doc,
             "All four players are SIMILAR-SIZED operators in the Spanish "
             "sustainable-camper niche. Approximate fleet size: Nomade < 30 vans, "
             "Further VAN ~ 50, People Camper ~ 60, Ocean Vans ~ 55. This is a "
             "PEER SET, not a David-vs-Goliath fight. The implication matters: "
             "the cost-structure advantage of the incumbents is LIMITED — their "
             "price premium is mostly driven by BRAND RECOGNITION and first-mover "
             "status, not by radically lower unit costs. Nomade can therefore "
             "close the gap by building brand perception on its stronger "
             "attributes (sustainability, customisation, design).",
             title='Size parity — the competitive set',
             fill='F5F3FF', title_color=RGBColor(0x5B, 0x21, 0xB6))

    # ---------- 2. Company summary ----------
    h(doc, '2. Company summary — Nomade vs. market')
    para(doc,
         "Nomade Vans is a NEW entrant in the campervan rental market, specialised "
         "in sustainable and highly customisable van designs with fast delivery "
         "(within 1 month). Despite its superior features it faces the classic "
         "challenge of low brand awareness while competing against well-established "
         "incumbents.")
    tbl = doc.add_table(rows=6, cols=3); tbl.style = 'Light Grid Accent 1'
    make_header_row(tbl, ['Attribute', 'Nomade', 'Market average'])
    for i, row in enumerate([
        ('Market position',     'NEW ENTRANT',  'Established (8–9/10)'),
        ('Sustainability',      '10/10',        '4–6/10'),
        ('Customisation',       '10/10',        '3–6/10'),
        ('Brand recognition',   '3/10',         '8–9/10'),
        ('Service quality',     '7/10',         '7–9/10'),
    ], 1):
        for j, v in enumerate(row):
            tbl.rows[i].cells[j].text = v
    doc.add_paragraph()

    # ---------- 2. Data note ----------
    note_box(doc,
             "Competitor prices are only published as DAILY tariffs. Weekend "
             "and weekly rates in the simulator are ESTIMATES aggregated from "
             "the daily tariff (industry convention: weekend ≈ 5–10 % per-day "
             "discount, week ≈ 20–30 % total discount). Treat them as reference "
             "scenarios. Nomade WTP data itself is real (95 respondents).",
             title='⚠ Data note — incomplete dataset (v3.0)')

    # ---------- 3. Data available ----------
    h(doc, '3. Data available')

    # 3.1 Standard survey
    h(doc, '3.1 Survey — Standard camper (WTP)', level=2)
    for lbl, prices, resps in [
        ('Daily rentals',            [50,60,70,80,90,100,110,120,130],  [17,14,10,16,12,19,2,4,1]),
        ('Weekend (2-night) rentals',[80,100,120,140,160,180,200],      [12,7,18,13,11,15,11]),
        ('Weekly (7-night) rentals', [240,300,360,420,480,540,600,660,720],[14,10,10,15,15,8,10,3,6]),
    ]:
        para(doc, lbl, bold=True)
        tbl = doc.add_table(rows=2, cols=len(prices)+1); tbl.style = 'Light Grid Accent 1'
        tbl.rows[0].cells[0].text = 'Price (€)'
        tbl.rows[1].cells[0].text = 'Responses'
        for i, p in enumerate(prices, 1):
            tbl.rows[0].cells[i].text = str(p)
            tbl.rows[1].cells[i].text = str(resps[i-1])
        doc.add_paragraph()

    # 3.2 Premium survey
    h(doc, '3.2 Survey — Premium camper (WTP)', level=2)
    for lbl, prices, resps in [
        ('Daily rentals',            [60,70,80,90,100,110,120,130,140,150,160,180,200], [10,5,12,8,14,12,19,5,3,4,1,1,1]),
        ('Weekend (2-night) rentals',[120,140,160,180,200,220,240,260,280],              [10,6,13,8,22,2,18,4,2]),
        ('Weekly (7-night) rentals', [360,420,480,540,600,660,720,780,840],              [14,7,4,13,17,8,13,7,3]),
    ]:
        para(doc, lbl, bold=True)
        tbl = doc.add_table(rows=2, cols=len(prices)+1); tbl.style = 'Light Grid Accent 1'
        tbl.rows[0].cells[0].text = 'Price (€)'
        tbl.rows[1].cells[0].text = 'Responses'
        for i, p in enumerate(prices, 1):
            tbl.rows[0].cells[i].text = str(p)
            tbl.rows[1].cells[i].text = str(resps[i-1])
        doc.add_paragraph()

    # 3.3 Cost structure
    h(doc, '3.3 Cost structure', level=2)
    tbl = doc.add_table(rows=4, cols=4); tbl.style = 'Light Grid Accent 1'
    make_header_row(tbl, ['Item', 'Daily', 'Weekend', 'Weekly'])
    for i, row in enumerate([
        ('Variable cost (€/day)',       '30', '30', '30'),
        ('Fixed cost (€/year)',         '69,750 (Std) / 77,750 (Prem)', '—', '—'),
        ('Average demand (rentals/yr)', '5,000', '1,825', '521.43'),
    ], 1):
        for j, v in enumerate(row):
            tbl.rows[i].cells[j].text = v
    doc.add_paragraph()

    # 3.4 Competitors (single-tariff, daily-only)
    h(doc, '3.4 Competitor data — single-tariff, daily-only', level=2)
    note_box(doc,
             'Important: the three competitors offer a SINGLE VERSION of their camper — '
             'there is NO Standard / Premium split on their side. They also publish '
             'DAILY TARIFFS ONLY. Any weekend / weekly number in the simulator is a '
             'reference projection from the daily tariff, not observed data. Nomade is '
             'the only player proposing two tiers (Standard + Premium) and three '
             'rental periods (Day / Weekend / Week).',
             title='Asymmetry in the competitive set',
             fill='E0F2FE', title_color=RGBColor(0x1E, 0x3A, 0x8A))

    tbl = doc.add_table(rows=5, cols=6); tbl.style = 'Light Grid Accent 1'
    make_header_row(tbl, ['Competitor','Daily tariff (€)','Tiers offered','Establishment 0–10','Brand 0–10','Sustainability 0–10'])
    for i, row in enumerate([
        ('Further VAN',   '85',           'One (single tier)', '8', '8', '5'),
        ('People Camper', '105',          'One (single tier)', '9', '9', '4'),
        ('Ocean Vans',    '98',           'One (single tier)', '9', '9', '6'),
        ('Nomade',        '80 / 100 (Std / Prem)', 'TWO (Standard + Premium)', '2', '3', '10'),
    ], 1):
        for j, v in enumerate(row):
            tbl.rows[i].cells[j].text = v
    doc.add_paragraph()
    para(doc, 'Competitor profiles:', bold=True)
    bullet(doc, 'Further VAN — well-established, affordable single-tier offer. Less differentiated; operational-efficiency focus.')
    bullet(doc, 'People Camper — premium single-tier positioning. High daily tariff suggests brand-premium pricing.')
    bullet(doc, 'Ocean Vans — sophisticated single-tier offer, mid-high daily price. Positioned on comfort + design.')
    bullet(doc, 'Nomade — NEW ENTRANT with two tiers. Highly differentiated (sustainable design, customisation, 1-month delivery) but low brand recognition. The two-tier structure alone is a differentiator vs the single-tier field.')

    # 3.5 Positioning maps
    h(doc, '3.5 Positioning maps — price vs. perceived attribute', level=2)
    para(doc,
         'A positioning map plots each player on two axes: price (Y) and a key '
         'perceived attribute (X). It is the most direct visual way to see where '
         'the market\'s white space is, where the closest competitor really sits, '
         'and why Nomade ends up charging less despite being objectively more '
         'differentiated.')

    para(doc, 'The four players\' coordinates (used by the simulator):', bold=True)
    tbl = doc.add_table(rows=5, cols=6); tbl.style = 'Light Grid Accent 1'
    make_header_row(tbl, ['Player','Daily price (€)','Brand 0–10','Sustainability 0–10','Design 0–10','Fleet size'])
    for i, row in enumerate([
        ('Nomade',        '80',  '3', '10', '9', '< 30 vans'),
        ('Further VAN',   '85',  '8', '5',  '7', '~ 50 vans'),
        ('People Camper', '105', '9', '4',  '6', '~ 60 vans'),
        ('Ocean Vans',    '98',  '9', '6',  '7', '~ 55 vans'),
    ], 1):
        for j, v in enumerate(row):
            tbl.rows[i].cells[j].text = v
    doc.add_paragraph()

    para(doc, 'Three useful maps — tick them on Competition-based → Graph picker:', bold=True)
    bullet(doc, 'Price × Brand recognition — Nomade is bottom-left (low brand, low price); competitors are top-right (high brand, high price). The price gap mirrors the brand gap almost perfectly.')
    bullet(doc, 'Price × Sustainability — the picture flips: Nomade is top-right (best score, 10/10) yet the LEAST expensive. The competitors sit bottom-left, with People Camper the weakest on sustainability (4/10). This map alone justifies a skimming move on Premium.')
    bullet(doc, 'Price × Design / customisation — Nomade is mid-top (9/10 on design, €80 price). Competitors cluster mid-price / mid-design. This shows that the pricing gap is NOT explained by product inferiority.')

    note_box(doc,
             'The three maps together tell the story: Nomade\'s lower price is NOT '
             'explained by inferior product or by a structural cost disadvantage '
             '(sizes are similar). It is explained almost entirely by the BRAND '
             'RECOGNITION gap. The strategic implication is therefore clear — '
             'invest in brand perception (narrative, channels, reviews, endorsements) '
             'to move UP the price axis without having to touch the product.',
             title='Why Nomade charges less — and what to do about it',
             fill='F5F3FF', title_color=RGBColor(0x5B, 0x21, 0xB6))

    para(doc,
         'A fourth axis — versatility (number of tariffs / tiers) — is where Nomade '
         'is ALONE: two tiers and three periods vs the competitors\' single-tier, '
         'daily-only offer. That asymmetry is itself a positioning lever '
         '(anchoring sandwich, flexible bundles) and is rendered in the simulator '
         'under the "Std+Prem × Day/Weekend/Week + competitors" chart.',
         italic=True, color=GREY)

    # ---------- 4. Tool features / simulator walkthrough ----------
    doc.add_page_break()
    h(doc, '4. Using the interactive simulator')
    para(doc, 'URL: ')
    p = doc.paragraphs[-1]; link = p.add_run(SIMULATOR_URL)
    link.font.color.rgb = BLUE; link.underline = True

    para(doc,
         'The simulator is a CLOSED LOOP. The Nomade price lives in one place and '
         'every tab reads or writes the same number. You do not have to re-type the '
         'same price in multiple tabs — pick one, iterate, and the Excel download '
         'will reflect whatever you left on "Position & prices".',
         italic=True, color=GREY)

    for s in [
        'Pick period (Day / Weekend / Week) and version (Standard / Premium) at the top.',
        'Competition-based → Profit curve: DRAG the green "Our price" line. Coloured dashed lines are competitors; KPIs (profit, demand, revenue, lost-vs-peak) update live. Use Match-competitor / Snap-to-peak quick-jump buttons.',
        'Psychological factors: all inputs pre-loaded from your current position. Experiment with Anchoring, Charm pricing, Prospect Theory (before/after), Reference price. Use Apply-to-Standard / Apply-to-Premium to push tested numbers into the final position.',
        'Position & prices: same interactive chart + numeric inputs for every period. Add the analyst note.',
        'Case answers: the four formal questions (see section 4) — write your answer in each box.',
        'Download Excel: single .xlsx with 8 sheets including Case answers and all 9 charts as images — this file IS the deliverable.',
    ]:
        numbered(doc, s)

    # ---------- 5. Strategic framework ----------
    doc.add_page_break()
    h(doc, '5. Strategic framework — the 4 scenarios')
    para(doc,
         'The pricing decision depends on comparing the WTP optimal price (Popt) '
         'with competitor prices (C) and on understanding whether the gap comes '
         'from cost structure or from differentiation.')
    tbl = doc.add_table(rows=9, cols=4); tbl.style = 'Light Grid Accent 1'
    make_header_row(tbl, ['Scenario', 'Status', 'Reason', 'Strategy'])
    for i, row in enumerate([
        ('a.1','Popt > C','More differentiated',     'NO — illogical'),
        ('a.2','Popt > C','Worse costs',             'YES (−) Penetration'),
        ('a.3','Popt < C','Less differentiated',     'YES (−−) Aggressive penetration'),
        ('a.4','Popt < C','Better costs',            'NO — great position'),
        ('b.1','Popt > C','More differentiated',     'RISKY — only for Apple-like brands'),
        ('b.2','Popt > C','Worse costs',             'NO — escape / reposition'),
        ('b.3','Popt < C','Less differentiated',     'YES (+) Skimming'),
        ('b.4','Popt < C','Better costs',            'NO — already right'),
    ], 1):
        for j, v in enumerate(row):
            tbl.rows[i].cells[j].text = v
    doc.add_paragraph()
    para(doc, 'Nomade\'s strategic position (v3 baseline):', bold=True)
    bullet(doc, 'Current status: Popt < C (€80 / €100 vs €85–€105 Std and €95–€129 Prem).')
    bullet(doc, 'Cost structure: similar or slightly better (sustainable-design efficiency).')
    bullet(doc, 'Differentiation: OBJECTIVELY HIGHER (sustainability 10/10, customisation 10/10).')
    bullet(doc, 'BUT low brand recognition (NEW entrant — 3/10 vs 8–9/10).')
    bullet(doc, 'This is SCENARIO b.3: Popt < C because customers perceive Nomade as less differentiated (the brand is unknown, not the product).')
    bullet(doc, 'Recommendation: SKIMMING (+) — moderate price increase to build brand perception and leverage the anchor effect.')

    # ---------- 6. FORMAL QUESTIONS (4 v3) ----------
    doc.add_page_break()
    h(doc, '6. Case questions')
    para(doc,
         'Four formal questions from Nomade v3.0. Write your answer either (a) '
         'inside the simulator\'s Case answers tab — which is exported into the '
         'Excel — or (b) directly in the answer boxes below.',
         italic=True)

    # Q1
    p = doc.add_paragraph()
    r = p.add_run('P1 — Competitive analysis and positioning')
    r.bold = True; r.font.color.rgb = ORANGE
    para(doc, 'With the data provided and the tool:', bold=True)
    bullet(doc, 'Give assumptions about the competitors\' cost structures and their brand-value differentiation.')
    bullet(doc, 'Remember NOMADE is NEW while the others are well-established.')
    bullet(doc, 'For each competitor, evaluate whether higher prices are due to: (a) higher operating costs, or (b) higher differentiation (brand recognition, quality, features).')
    bullet(doc, 'Build a positioning map of Nomade vs. competitors on the key attributes.')
    para(doc, 'Simulator tasks:', bold=True, color=GREY)
    bullet(doc, 'Read profit at each competitor\'s dashed line on the Profit curve.')
    bullet(doc, 'Tick the "Std+Prem × Day/Weekend/Week + competitors" chart on Competition-based to see the 2-tier vs 1-tier asymmetry at a glance.')
    note_box(doc,
             '(a) COST STRUCTURE — Established players (Further VAN, People Camper, '
             'Ocean Vans) enjoy amortised fleets, procurement scale and branded '
             'channels ⇒ LOWER variable cost per rental but HIGHER fixed marketing / '
             'brand costs. Nomade is new ⇒ SIMILAR variable cost (sustainable-design '
             'efficiency compensates scale), HIGHER fixed cost per unit (thin fleet). '
             '(b) DIFFERENTIATION — their price premium is explained mostly by BRAND '
             'RECOGNITION, not by objectively superior product features. '
             '(c) POSITIONING MAP — Nomade (Establishment 2/10, Brand 3/10, '
             'Sustainability 10/10, Design 9/10) vs Further VAN (8/8/5/7), People '
             'Camper (9/9/4/6), Ocean Vans (9/9/6/7). Nomade is OBJECTIVELY MORE '
             'DIFFERENTIATED but PERCEIVED AS LESS so — brand is the gap to close.',
             title='Answer guide — P1',
             fill='ECFDF5', title_color=RGBColor(0x06, 0x5F, 0x46))
    para(doc, 'Your answer:', bold=True)
    answer_box(doc, lines=12)

    # Q2
    p = doc.add_paragraph()
    r = p.add_run('P2 — Apply the 4-scenario strategic framework')
    r.bold = True; r.font.color.rgb = ORANGE
    para(doc, 'Using the 4-scenario framework (section 4):', bold=True)
    bullet(doc, 'Identify which scenario applies to Nomade for each version (Standard / Premium).')
    bullet(doc, 'Is Popt > or < competitors\' prices? Is the reason cost structure or differentiation?')
    bullet(doc, 'Given the scenario, should Nomade pursue: (a) penetration (−) / aggressive penetration (−−), (b) skimming (+) / fast-aggressive (++), or (c) nothing?')
    bullet(doc, 'Justify with competition theory and the WTP profit analysis.')
    note_box(doc,
             'WTP optima: Standard €80 · Premium €100. Competitor range (daily) '
             '€85–€105. Therefore Popt < C. Cost structure similar or slightly '
             'better (sustainable-design efficiency); differentiation OBJECTIVELY '
             'HIGHER but BRAND RECOGNITION LOW. This matches SCENARIO b.3 (Popt < C '
             'because perceived as less differentiated, despite real superiority). '
             'RECOMMENDATION → SKIMMING (+) on the Premium tier (raise gradually to '
             'build brand anchor) and a gentle PENETRATION (−) on Standard only for '
             'the first season (seed trials). Never (−−) aggressive penetration — it '
             'burns margin without fixing the brand gap. Confirm the decision on the '
             'Profit curve: check the profit at +€5 and +€10 before committing.',
             title='Answer guide — P2',
             fill='ECFDF5', title_color=RGBColor(0x06, 0x5F, 0x46))
    para(doc, 'Your answer:', bold=True)
    answer_box(doc, lines=14)

    # Q3
    p = doc.add_paragraph()
    r = p.add_run('P3 — Weekend / week pricing strategy')
    r.bold = True; r.font.color.rgb = ORANGE
    para(doc, 'Competitor data is DAILY only:', bold=True)
    bullet(doc, 'Start from your daily pricing decisions (WTP optimum + competitive positioning).')
    bullet(doc, 'Assume industry conventions: weekend ≈ 5–10 % per-day discount, week ≈ 20–30 % total discount.')
    bullet(doc, 'Propose Nomade\'s weekend / week prices.')
    bullet(doc, 'Use anchoring (daily × 7 vs weekly) and the "sandwich" between Standard and Premium.')
    note_box(doc,
             'Use weekend ≈ −5 % per day (2-night packages) and week ≈ −30 % per day '
             '(long-stay discount). Proposed Nomade table: Std (Day €80 / Weekend '
             '€75 per day / Week €60 per day) · Prem (Day €100 / Weekend €95 per '
             'day / Week €77 per day). Anchor the week price by displaying the '
             'per-day saving vs the "daily × 7" reference. On Psychological '
             'factors → Anchoring, keep Basic = cheapest competitor, Middle = '
             'Nomade Std, Premium = Nomade Prem — click Apply-to-position. Use the '
             '"Std+Prem × Day/Weekend/Week + competitors" chart to verify the 2×3 '
             'sandwich visually.',
             title='Answer guide — P3',
             fill='ECFDF5', title_color=RGBColor(0x06, 0x5F, 0x46))
    para(doc, 'Your answer:', bold=True)
    answer_box(doc, lines=14)

    # Q4
    p = doc.add_paragraph()
    r = p.add_run('P4 — Complete pricing proposal')
    r.bold = True; r.font.color.rgb = ORANGE
    para(doc, 'Deliver the full pricing strategy with justification:', bold=True)
    bullet(doc, 'Complete price table (Standard / Premium × Day / Weekend / Week).')
    tbl = doc.add_table(rows=3, cols=4); tbl.style = 'Light Grid Accent 1'
    make_header_row(tbl, ['', 'Day', 'Weekend (per day)', 'Week (per day)'])
    tbl.rows[1].cells[0].text = 'Standard'
    tbl.rows[2].cells[0].text = 'Premium'
    doc.add_paragraph()
    bullet(doc, 'Comparison vs competitors (with your weekend/week assumptions).')
    bullet(doc, 'Justify with: (a) WTP analysis (profit curves), (b) competitive positioning (differentiation vs costs), (c) psychological tactics (anchoring, sandwich, charm).')
    bullet(doc, 'Expected business outcomes (profit, market positioning).')
    bullet(doc, 'Consider alternative structures (e.g., 3-tier with "enhanced" version).')
    bullet(doc, 'Risk mitigation (competitor reactions).')
    note_box(doc,
             'Baseline proposal (with charm endings): Standard Day €79 / Weekend '
             '€75/day / Week €60/day · Premium Day €99 / Weekend €95/day / Week '
             '€77/day. Charm €79 / €99 captures the digit-9 uplift without eroding '
             'margin. Expected outcome: ≈ 65 % acceptance on Std, 40–50 % on Prem, '
             '−€5 to −€10 gap vs competitor average, total profit > €80 k/yr '
             'combined. Alternative 3-tier (Base / Std / Premium) would deepen the '
             'anchoring effect but adds fleet complexity — defer to year 2. Risk '
             'mitigation: (i) competitors undercut daily Std → 6-month price lock + '
             'loyalty week-at-day-price; (ii) brand-recognition handicap → push '
             'sustainability narrative on every touchpoint; (iii) cannibalisation → '
             'keep the Std–Prem gap ≥ €20. Fill every price on Position & prices, '
             'add the analyst note, download the Excel — that file is the proposal.',
             title='Answer guide — P4',
             fill='ECFDF5', title_color=RGBColor(0x06, 0x5F, 0x46))
    para(doc, 'Your answer:', bold=True)
    answer_box(doc, lines=16)

    # ---------- 7. Evaluation ----------
    doc.add_page_break()
    h(doc, '7. Evaluation criteria')
    bullet(doc, 'Rigour — evidence inside the simulator (drag results, screenshots, numbers).')
    bullet(doc, 'Integration of competition + WTP + psychological factors into a single recommendation.')
    bullet(doc, 'Correct use of the 4-scenario strategic framework.')
    bullet(doc, 'Clarity and brevity — numbers justified, trade-offs explicit.')

    h(doc, 'Deliverable', level=2)
    para(doc,
         'The Excel (.xlsx) downloaded from the simulator is enough. Submit that file alone — '
         'it already contains every piece of evidence required:')
    bullet(doc, 'Guide — case background, data note and active view.')
    bullet(doc, 'Questions — the 13 survey items (reference only).')
    bullet(doc, 'Answers — Day / Weekend / Week — aggregated WTP distributions.')
    bullet(doc, 'Position — your final prices, competitor prices, peak / gap / acceptance KPIs and the analyst note.')
    bullet(doc, 'Case answers — your written answers to the four formal questions.')
    bullet(doc, 'Charts — the nine simulator charts embedded as images.')
    para(doc,
         'Before downloading, make sure you have filled the four Case answers and set the '
         'final prices in Position & prices. No additional Word or PowerPoint file is required.',
         italic=True, color=GREY)

    # ---------- APPENDIX ----------
    doc.add_page_break()
    h(doc, 'Appendix — The 13 survey questions (reference only)')
    para(doc,
         'These are the 13 items asked to the 95 survey respondents. They appear here only as '
         'reference so that you can interpret the aggregated WTP distribution displayed by the '
         'simulator. You are NOT asked to answer them yourself. Original Spanish wording in italics.',
         italic=True, color=GREY)

    items = [
        ('Q1',  'Please indicate your gender.',
                'Por favor, indique su género.'),
        ('Q2',  'What age range are you in?  <20 / 20-30 / 30-40 / 40-50 / 50-60 / >60',
                '¿En qué rango de edad se encuentra?'),
        ('Q3',  'What is your civil status?  Single / Married / Divorced / Widowed',
                '¿Cuál es su estado civil?'),
        ('Q4',  'What is your level of education?  Primary / Secondary / University / Master or PhD',
                '¿Cuál es su nivel de educación?'),
        ('Q5',  'Select your employment status.  Employed / Self-employed / Unemployed / Retired / Student',
                'Seleccione su situación laboral.'),
        ('Q6',  'What do you like to do in your free time? (multiple choice)  Travel · Reading · Sports · Adventure · Family time · Other',
                '¿Qué le gusta hacer en su tiempo libre?'),
        ('Q7',  'How many days do you usually go on vacation?  <3 / 3-5 / 5-7 / >7',
                '¿Cuántos días sueles irte de vacaciones?'),
        ('Q8',  'Maximum price you are willing to pay — STANDARD camper, per day.',
                'Precio máximo por el alquiler de una camper ESTÁNDAR al día.'),
        ('Q9',  'Maximum price you are willing to pay — STANDARD camper, weekend (2 nights).',
                'Precio máximo — camper ESTÁNDAR fin de semana (2 noches).'),
        ('Q10', 'Maximum price you are willing to pay — STANDARD camper, week (6 nights).',
                'Precio máximo — camper ESTÁNDAR semana (6 noches).'),
        ('Q11', 'Maximum price you are willing to pay — PREMIUM camper, per day.',
                'Precio máximo — camper PREMIUM al día.'),
        ('Q12', 'Maximum price you are willing to pay — PREMIUM camper, weekend (2 nights).',
                'Precio máximo — camper PREMIUM fin de semana (2 noches).'),
        ('Q13', 'Maximum price you are willing to pay — PREMIUM camper, week (6 nights).',
                'Precio máximo — camper PREMIUM semana (6 noches).'),
    ]
    for qid, english, spanish in items:
        p = doc.add_paragraph()
        run = p.add_run(f'{qid}. '); run.bold = True
        p.add_run(english)
        p2 = doc.add_paragraph()
        r2 = p2.add_run(spanish); r2.italic = True; r2.font.color.rgb = GREY; r2.font.size = Pt(10)

    # ---------- Save ----------
    stamp = datetime.now().strftime('%Y-%m-%d')
    out = f'Nomade_Vans_Case_v3.0_Simulator_Edition_{stamp}.docx'
    doc.save(out)
    print('Saved:', out)


if __name__ == '__main__':
    build()
